import streamlit as st
from PIL import Image
import pandas as pd
from pathlib import Path
import zipfile
import io
from datetime import datetime
from docx import Document
from docx.shared import RGBColor

from src.data_loader import SROIEDataLoader
from src.ocr_engine import ConfidenceOCR
from src.confidence_scorer import ConfidenceAnalyzer
from src.visualizer import ConfidenceVisualizer
from src.evaluator import OCREvaluator


# Page config
st.set_page_config(
    page_title="Confidence-Aware OCR",
    page_icon="📄",
    layout="wide"
)


# Initialize components (cache for performance)
@st.cache_resource
def init_components():
    """Initialize OCR components (cached)"""
    ocr = ConfidenceOCR()
    analyzer = ConfidenceAnalyzer()
    visualizer = ConfidenceVisualizer()
    evaluator = OCREvaluator()
    return ocr, analyzer, visualizer, evaluator


@st.cache_resource
def load_data_loader():
    """Load dataset (cached)"""
    try:
        loader = SROIEDataLoader("data/raw/SROIE2019")
        return loader
    except:
        return None


def process_single_image(image, filename, ocr, analyzer, visualizer):
    """Process single image through pipeline"""
    
    # Run OCR
    ocr_result = ocr.extract_text_with_confidence(image)
    
    # Analyze confidence
    analysis = analyzer.analyze(ocr_result)
    
    # Create visualization
    confidence_map = visualizer.visualize_confidence(image, ocr_result)
    review_map = visualizer.highlight_review_areas(image, analysis)
    
    # Create summary
    summary = {
        'filename': filename,
        'total_words': analysis['statistics']['total_words'],
        'avg_confidence': analysis['statistics']['avg_confidence'],
        'high_conf_count': analysis['high_confidence']['count'],
        'medium_conf_count': analysis['medium_confidence']['count'],
        'low_conf_count': analysis['low_confidence']['count'],
        'review_percentage': analysis['statistics']['review_percentage'],
        'needs_review': analysis['statistics']['needs_review']
    }
    
    return {
        'ocr_result': ocr_result,
        'analysis': analysis,
        'confidence_map': confidence_map,
        'review_map': review_map,
        'summary': summary
    }


def process_batch(uploaded_files, ocr, analyzer, visualizer):
    """Process multiple images"""
    
    results = []
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for idx, uploaded_file in enumerate(uploaded_files):
        status_text.text(f"Processing {idx + 1}/{len(uploaded_files)}: {uploaded_file.name}")
        
        try:
            image = Image.open(uploaded_file)
            result = process_single_image(
                image, 
                uploaded_file.name,
                ocr, analyzer, visualizer
            )
            results.append(result)
        except Exception as e:
            st.warning(f"⚠️ Error processing {uploaded_file.name}: {e}")
        
        progress_bar.progress((idx + 1) / len(uploaded_files))
    
    status_text.text("✅ Batch processing complete!")
    return results


def create_word_doc(result):
    """Create Word document with colored text based on confidence"""
    
    doc = Document()
    
    # Add title
    doc.add_heading(f"OCR Results: {result['summary']['filename']}", level=1)
    
    # Add metadata
    doc.add_paragraph(f"Total Words: {result['summary']['total_words']}")
    doc.add_paragraph(f"Average Confidence: {result['summary']['avg_confidence']:.1f}%")
    doc.add_paragraph(f"Words Needing Review: {result['summary']['low_conf_count']}")
    doc.add_paragraph("")
    
    # Add heading for extracted text
    doc.add_heading("Extracted Text (Color-coded by confidence)", level=2)
    
    # Legend
    legend = doc.add_paragraph()
    legend.add_run("Legend: ").bold = True
    
    high_run = legend.add_run("Green = High confidence (≥80%)  ")
    high_run.font.color.rgb = RGBColor(0, 128, 0)
    
    medium_run = legend.add_run("Orange = Medium confidence (60-79%)  ")
    medium_run.font.color.rgb = RGBColor(255, 140, 0)
    
    low_run = legend.add_run("Red = Low confidence (<60%)")
    low_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_paragraph("")
    
    # Add extracted text with colors
    paragraph = doc.add_paragraph()
    
    for word_info in result['ocr_result']['words']:
        text = word_info['text']
        confidence = word_info['confidence']
        
        # Add word with appropriate color
        run = paragraph.add_run(text + " ")
        
        if confidence >= 80:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif confidence >= 60:
            run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            run.bold = True  # Bold for low confidence
    
    return doc


def create_word_docs_zip(results):
    """Create ZIP file with Word documents for all results"""
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for result in results:
            # Create Word doc
            doc = create_word_doc(result)
            
            # Save to buffer
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Add to zip
            filename = result['summary']['filename'].replace('.jpg', '.docx').replace('.jpeg', '.docx').replace('.png', '.docx')
            zip_file.writestr(f"word_docs/{filename}", doc_buffer.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer


def create_images_zip(results):
    """Create zip file with all annotated images and CSV"""
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Add confidence maps
        for result in results:
            img_buffer = io.BytesIO()
            img = result['confidence_map']
            
            # Convert RGBA to RGB if needed
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            img.save(img_buffer, format='JPEG')
            zip_file.writestr(
                f"confidence_maps/{result['summary']['filename']}",
                img_buffer.getvalue()
            )
        
        # Add review maps
        for result in results:
            img_buffer = io.BytesIO()
            img = result['review_map']
            
            # Convert RGBA to RGB if needed
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            img.save(img_buffer, format='JPEG')
            zip_file.writestr(
                f"review_maps/{result['summary']['filename']}",
                img_buffer.getvalue()
            )
        
        # Add CSV summary
        summaries = [r['summary'] for r in results]
        df = pd.DataFrame(summaries)
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False)
        zip_file.writestr('batch_results_summary.csv', csv_buffer.getvalue())
    
    zip_buffer.seek(0)
    return zip_buffer


def main():
    """Main Streamlit app"""
    
    # Title
    st.title(" CONFIDENCE AWARE OCR Document Reader")
    st.markdown("Extract text with confidence scoring • Single & Batch Processing")
    
    # Sidebar
    st.sidebar.header("Settings")
    
    # Confidence thresholds
    st.sidebar.subheader("Confidence Thresholds")
    low_threshold = st.sidebar.slider("Low Confidence", 0, 100, 60)
    high_threshold = st.sidebar.slider("High Confidence", 0, 100, 80)
    
    # Initialize components
    ocr, analyzer, visualizer, evaluator = init_components()
    analyzer.low_threshold = low_threshold
    analyzer.high_threshold = high_threshold
    visualizer.low_threshold = low_threshold
    visualizer.high_threshold = high_threshold
    
    # Tabs
    tab1, tab2, tab3, tab4 = st.tabs([
        "Single Upload", 
        "Batch Processing", 
        "Dataset Explorer", 
        "About"
    ])
    
    # ========================================================================
    # TAB 1: Single Upload
    # ========================================================================
    with tab1:
        st.header("Upload Single Document")
        
        uploaded_file = st.file_uploader(
            "Choose an image file",
            type=['jpg', 'jpeg', 'png'],
            help="Upload a receipt or document image"
        )
        
        if uploaded_file:
            # Load image
            image = Image.open(uploaded_file)
            
            # Process
            with st.spinner("Processing document..."):
                result = process_single_image(
                    image, uploaded_file.name, ocr, analyzer, visualizer
                )
            
            ocr_result = result['ocr_result']
            analysis = result['analysis']
            confidence_map = result['confidence_map']
            review_map = result['review_map']
            
            # Display results
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader(" Original Image")
                st.image(image, use_container_width=True)
            
            with col2:
                st.subheader("Confidence Map")
                st.image(confidence_map, use_container_width=True)
                st.caption("🟢 Green = High | 🟡 Yellow = Medium | 🔴 Red = Low")
            
            # Statistics
            st.markdown("---")
            st.subheader("Analysis Results")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Total Words", analysis['statistics']['total_words'])
            
            with col2:
                st.metric("Avg Confidence", f"{analysis['statistics']['avg_confidence']:.1f}%")
            
            with col3:
                st.metric(
                    "Words Needing Review",
                    analysis['low_confidence']['count'],
                    delta=f"-{analysis['statistics']['review_percentage']:.1f}%",
                    delta_color="inverse"
                )
            
            with col4:
                st.metric("High Confidence", analysis['high_confidence']['count'])
            
            # Confidence distribution
            st.markdown("---")
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.subheader("Confidence Distribution")
                
                dist_data = pd.DataFrame({
                    'Category': ['High (≥80%)', 'Medium (60-79%)', 'Low (<60%)'],
                    'Count': [
                        analysis['high_confidence']['count'],
                        analysis['medium_confidence']['count'],
                        analysis['low_confidence']['count']
                    ]
                })
                
                st.bar_chart(dist_data.set_index('Category')['Count'])
            
            with col2:
                st.subheader("Summary")
                st.dataframe(dist_data, hide_index=True)
            
            # Extracted text
            st.markdown("---")
            st.subheader("Extracted Text")
            st.text_area("Full OCR Output", ocr_result['full_text'], height=200)
            
            # Low confidence words
            if analysis['low_confidence']['count'] > 0:
                st.markdown("---")
                st.subheader("Low Confidence Words (Needs Review)")
                
                low_conf_df = pd.DataFrame([
                    {'Text': w['text'], 'Confidence': f"{w['confidence']}%"}
                    for w in analysis['low_confidence']['words']
                ])
                
                st.dataframe(low_conf_df, hide_index=True)
                
                st.subheader("🔴 Review Map")
                st.image(review_map, use_container_width=True)
    
    # ========================================================================
    # TAB 2: Batch Processing
    # ========================================================================
    with tab2:
        st.header("Batch Processing")
        st.markdown("Upload multiple documents to process them all at once")
        
        uploaded_files = st.file_uploader(
            "Choose multiple image files",
            type=['jpg', 'jpeg', 'png'],
            accept_multiple_files=True,
            help="Select multiple receipt/document images"
        )
        
        if uploaded_files and len(uploaded_files) > 0:
            st.info(f" {len(uploaded_files)} files uploaded")
            
            # Smart limit warning
            if len(uploaded_files) > 50:
                st.warning("⚠️ **Performance Notice**: Processing >50 documents may be slow in the web interface.")
                st.info("💡 **Tip**: For large batches (100+), use the CLI: `python main.py`")
                
                proceed = st.checkbox("⚡ Process anyway (may take several minutes)")
                if not proceed:
                    st.stop()
            
            if st.button(" Process All Documents", type="primary"):
                
                # Process batch
                with st.spinner(f"Processing {len(uploaded_files)} documents..."):
                    results = process_batch(uploaded_files, ocr, analyzer, visualizer)
                
                if not results:
                    st.error("No documents were successfully processed")
                    return
                
                # Store results in session state for downloads
                st.session_state['batch_results'] = results
                
                # Extract summaries
                summaries_df = pd.DataFrame([r['summary'] for r in results])
                
                # Display aggregate statistics
                st.markdown("---")
                st.subheader("Batch Processing Results")
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("Total Documents", len(results))
                
                with col2:
                    st.metric(
                        "Avg Confidence",
                        f"{summaries_df['avg_confidence'].mean():.1f}%"
                    )
                
                with col3:
                    needs_review = summaries_df['needs_review'].sum()
                    st.metric("Needs Review", needs_review)
                
                with col4:
                    st.metric(
                        "Total Words",
                        int(summaries_df['total_words'].sum())
                    )
                
                # Distribution chart
                st.markdown("---")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("Confidence Distribution Across All Documents")
                    
                    agg_data = pd.DataFrame({
                        'Category': ['High', 'Medium', 'Low'],
                        'Total Count': [
                            summaries_df['high_conf_count'].sum(),
                            summaries_df['medium_conf_count'].sum(),
                            summaries_df['low_conf_count'].sum()
                        ]
                    })
                    
                    st.bar_chart(agg_data.set_index('Category')['Total Count'])
                
                with col2:
                    st.subheader("Aggregate Summary")
                    st.dataframe(agg_data, hide_index=True)
                
                # Detailed results table
                st.markdown("---")
                st.subheader("📄 Detailed Results for Each Document")
                
                # Format the dataframe for display
                display_df = summaries_df[[
                    'filename', 'total_words', 'avg_confidence', 
                    'low_conf_count', 'review_percentage'
                ]].copy()
                
                display_df.columns = [
                    'Filename', 'Words', 'Avg Conf %', 
                    'Low Conf Words', 'Review %'
                ]
                
                # Highlight rows needing review
                def highlight_review(row):
                    if row['Review %'] > 30:
                        return ['background-color: #ffcccc'] * len(row)
                    elif row['Review %'] > 20:
                        return ['background-color: #fff4cc'] * len(row)
                    else:
                        return [''] * len(row)
                
                styled_df = display_df.style.apply(highlight_review, axis=1)
                st.dataframe(styled_df, hide_index=True, use_container_width=True)
                
                st.caption("🔴 Red highlight = >30% needs review | 🟡 Yellow = >20% needs review")
                
                # Download options
                st.markdown("---")
                st.subheader("Download Results")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    # Download CSV
                    csv_data = summaries_df.to_csv(index=False)
                    st.download_button(
                        label="📊 CSV Summary",
                        data=csv_data,
                        file_name=f"batch_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                        mime="text/csv",
                        use_container_width=True,
                        help="Download summary statistics as CSV"
                    )
                
                with col2:
                    # Download annotated images ZIP
                    with st.spinner("Preparing images ZIP..."):
                        images_zip = create_images_zip(results)
                    
                    st.download_button(
                        label="📦 ZIP: Images + CSV",
                        data=images_zip,
                        file_name=f"ocr_annotated_images_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip",
                        use_container_width=True,
                        help="Download all annotated images + CSV"
                    )
                
                with col3:
                    # Download Word docs ZIP
                    with st.spinner("Creating Word documents..."):
                        word_zip = create_word_docs_zip(results)
                    
                    st.download_button(
                        label="📝 ZIP: Word Docs",
                        data=word_zip,
                        file_name=f"ocr_word_docs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip",
                        use_container_width=True,
                        help="Download color-coded Word documents"
                    )
                
                st.info("💡 **Word Docs**: Green = High confidence, Orange = Medium, Red (Bold) = Low confidence (needs review)")
                
                # Show sample images
                st.markdown("---")
                st.subheader("Sample Results")
                
                # Show first 3 documents
                num_samples = min(3, len(results))
                cols = st.columns(num_samples)
                
                for idx, col in enumerate(cols):
                    with col:
                        st.image(
                            results[idx]['confidence_map'],
                            caption=f"{results[idx]['summary']['filename']}",
                            use_container_width=True
                        )
                        st.caption(f"Confidence: {results[idx]['summary']['avg_confidence']:.1f}%")
        else:
            st.info("👆 Upload multiple documents to get started with batch processing")
            st.markdown("**💡 Tips:**")
            st.markdown("- Upload 5-10 images for quick demo")
            st.markdown("- Upload up to 50 for comprehensive analysis")
            st.markdown("- For 100+ documents, use CLI: `python main.py`")
    
    # ========================================================================
    # TAB 3: Dataset Explorer
    # ========================================================================
    with tab3:
        st.header("Explore SROIE Dataset")
        
        loader = load_data_loader()
        
        if loader is None:
            st.warning("⚠️ SROIE dataset not found. Please check data/raw/SROIE2019/")
        else:
            samples = loader.get_sample_ids('train')
            
            st.success(f"Found {len(samples)} receipts in dataset")
            
            selected_sample = st.selectbox("Select a receipt to analyze", samples, index=0)
            
            if st.button("Analyze Selected Receipt"):
                sample = loader.load_complete_sample(selected_sample)
                
                result = process_single_image(
                    sample['image'], selected_sample, ocr, analyzer, visualizer
                )
                
                report = evaluator.generate_report(
                    selected_sample,
                    result['ocr_result'],
                    result['analysis'],
                    sample['boxes'],
                    sample['entities']
                )
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("Receipt")
                    st.image(sample['image'], use_container_width=True)
                
                with col2:
                    st.subheader("Confidence Map")
                    st.image(result['confidence_map'], use_container_width=True)
                
                st.markdown("---")
                st.subheader("Evaluation Results")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("Text Similarity", f"{report['text_similarity']:.1f}%")
                
                with col2:
                    st.metric("Field Accuracy", f"{report['field_accuracy']:.1f}%")
                
                with col3:
                    st.metric("Avg Confidence", f"{report['avg_confidence']:.1f}%")
                
                st.markdown("---")
                st.subheader("✅ Ground Truth Entities")
                
                entities_df = pd.DataFrame([
                    {'Field': k, 'Value': v}
                    for k, v in sample['entities'].items()
                ])
                
                st.dataframe(entities_df, hide_index=True)
    
    # ========================================================================
    # TAB 4: About
    # ========================================================================
    with tab4:
        st.header("About This Project")
        
        st.markdown("""
                    ## ℹ️ About This System
Even in 2026, organizations still rely heavily on scanned receipts, invoices, and paper-based forms.  
OCR technology exists  but the reality in business environments is harsh:

- OCR output is **rarely perfect**
- Errors in critical fields (amounts, dates, vendor names) can trigger **financial loss, compliance risk, or operational delays**
- Existing solutions either **hide uncertainty** or force humans to **review everything manually**

This creates a costly gap between automation and trust.

The **Confidence Aware OCR Document Reader** is designed to close that gap.

Instead of asking *“Can OCR read this?”*, the system asks a more practical business question:

**“Which parts of this OCR output can be trusted — and which parts require human attention?”**

Rather than chasing unrealistic 100% automation, this system makes **uncertainty visible, measurable, and actionable**.



## 🔑 How This System Is Different

- **Confidence-Aware OCR**  
  Every extracted word is assigned a confidence score (0–100%), indicating how reliable the OCR result is.

- **Visual Review Maps**  
  Text is color coded directly on the document:
  - 🟢 Green: High confidence — usually correct  
  - 🟡 Yellow: Medium confidence — review if important  
  - 🔴 Red: Low confidence — requires human verification

- **Human in the Loop Efficiency**  
  Reviewers no longer check entire documents.  
  They focus only on uncertain regions typically **20–30% of the text** — saving time without increasing risk.

- **Validated on Real-World Data**  
  Performance is evaluated using the **SROIE2019 dataset**, which contains noisy, real-world scanned receipts with varied layouts and distortions.



## 💡 Real-World Business Impact

### Save Time & Reduce Manual Effort  
Most teams review every scanned document end-to-end because OCR gives no signal of reliability.  
This system highlights only what matters — reducing manual review effort by **up to 70–80%**.

###  Minimize Costly Errors  
Mistakes in invoice totals, tax amounts, or customer data can lead to financial losses or audit issues.  
Low-confidence text is clearly flagged, ensuring critical errors are caught **before** downstream impact.

### Improve Operational Throughput  
Faster document processing means **more volume handled without increasing headcount**.  
Teams can scale operations while maintaining accuracy and control.

### Enable Smarter Oversight  
Managers, auditors, and QA teams instantly see **where human attention is required**, instead of relying on blind trust or excessive rechecking.

###  Applicable Across Industries  
Any workflow involving scanned documents benefits:
- Finance & Accounting
- Legal & Compliance
- Healthcare Records
- Logistics & Operations

The system adapts without changing existing document processes.



## 🏗️ Technology Overview

- **OCR Engine:** Tesseract OCR (word-level confidence extraction)
- **Processing Stack:** Python, PIL, NumPy, pandas
- **Web Interface:** Streamlit with single & batch processing
- **Exports:** CSV summaries, annotated images, color-coded Word documents
- **Dataset:** SROIE2019 (real-world scanned receipts)



## Why This Matters

This project demonstrates how AI can be applied **responsibly and practically** —  
not by replacing humans, but by **directing human effort where it matters most**.

It transforms OCR from a fragile automation step into a **decision-support system**.

        
        
### 📊 Performance Metrics
- Average confidence: ~70-75%
- High confidence accuracy: ~94%
- Low confidence accuracy: ~40%
- **Time saved**: Review only 20-30% of words instead of 100%
- **Batch processing**: 10-50 documents in minutes
        
        
### 👨‍💻 Author
        
**Shapna**  
Machine Learning Engineer Intern Project  

        
📧 shapna3006@gmail.com
        
         
**Demonstrates**: OCR, CV, Batch Processing, Production Thinking
        """)
        
        
    
if __name__ == "__main__":
    main()